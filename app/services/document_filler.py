"""
Модуль «Документы» — генерация официальных бланков госорганов в формате .docx.

Принцип честности: мы автоматически заполняем ТОЛЬКО то, что реально знаем
(регион, дату, правовое основание из уже сгенерированного текста обращения).
Персональные поля (ФИО, СНИЛС, адрес, паспорт) оставляем пустыми — сервис
анонимный и не хранит и не запрашивает персональные данные пользователя.
Пользователь дозаполняет их от руки или в Word перед подачей.

Первый реализованный бланк: «Заявление о перерасчёте размера пенсии»
(Приложение N 2 к Административному регламенту, утв. Постановлением
Правления ПФ РФ от 23.01.2019 N 16п, ред. от 23.09.2020).

Второй бланк: «Заявление о назначении ежемесячной денежной выплаты» (ЕДВ)
(Приложение N 1 к Административному регламенту, утв. Постановлением
Правления ПФ РФ от 19.08.2019 N 414п, ред. от 23.09.2020).
"""

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.data.region_declension import to_dative

# Ключи (категория, подкатегория), для которых уже реализовано заполнение бланка.
SUPPORTED_DOCUMENTS = {
    ("pension", "pension_recalculation"): "pension_recalculation_form",
    ("pension", "pension_underpayment"): "pension_recalculation_form",
    ("health", "disability"): "edv_form",
    ("health", "veteran"): "edv_form",
}

# Правовое основание ЕДВ по подкатегории — заполняем только то, что реально
# знаем по выбору пользователя; остальные поля (в т.ч. точная "категория"
# из п.4 формы, т.к. внутри "инвалидности"/"ветеранства" много подвариантов)
# оставляем пустыми для самостоятельного заполнения.
EDV_LEGAL_BASIS = {
    "disability": "Федеральным законом от 24.11.1995 № 181-ФЗ «О социальной защите инвалидов в Российской Федерации»",
    "veteran": "Федеральным законом от 12.01.1995 № 5-ФЗ «О ветеранах»",
}


def is_document_supported(category: str, subcategory: str) -> bool:
    return (category, subcategory) in SUPPORTED_DOCUMENTS


def _add_blank_line(doc: Document, label: str, width_chars: int = 60):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: " + "_" * width_chars)
    run.font.size = Pt(11)
    return p


def generate_pension_recalculation_form(region_name: str, reason_text: str) -> bytes:
    """
    Генерирует заполненный (частично) бланк «Заявление о перерасчёте размера
    пенсии» строго по официальной форме (Приложение N 2, Постановление
    Правления ПФ РФ от 23.01.2019 N 16п).
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Наименование территориального органа — единственное автозаполненное поле шапки
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"В Отделение Фонда пенсионного и социального страхования\nРоссийской Федерации по {to_dative(region_name)}")
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ЗАЯВЛЕНИЕ\nО ПЕРЕРАСЧЕТЕ РАЗМЕРА ПЕНСИИ")
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()

    doc.add_paragraph(
        "1. _________________________________________________________________________"
    )
    doc.add_paragraph("(фамилия, имя, отчество (при наличии))").italic = True
    _add_blank_line(doc, "Страховой номер индивидуального лицевого счета (СНИЛС)", 30)
    _add_blank_line(doc, "Адрес места жительства", 55)
    _add_blank_line(doc, "Номер телефона", 25)

    doc.add_paragraph()
    doc.add_paragraph(
        "2. Представитель (если заявление подаётся представителем) — заполняется при необходимости, "
        "иначе оставить пустым."
    )

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("3. Прошу произвести перерасчёт размера ").bold = False
    p3.add_run("____________________________ ")
    run_note = p3.add_run("(вид пенсии — заполнить самостоятельно)")
    run_note.italic = True
    run_note.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("по следующему основанию (отметить нужное или использовать поле «иное»):")

    reasons = [
        "увеличение величины индивидуального пенсионного коэффициента за периоды до 1 января 2015 года",
        "увеличение суммы коэффициентов за иные периоды, засчитываемые в страховой стаж, после 1 января 2015 года",
        "наличие (увеличение количества) нетрудоспособных членов семьи, находящихся на иждивении пенсионера",
        "приобретение необходимого стажа работы в районах Крайнего Севера и приравненных местностях",
    ]
    for reason in reasons:
        doc.add_paragraph(f"☐  {reason}", style=None)

    p_other = doc.add_paragraph()
    p_other.add_run("☑  иное: ").bold = True
    p_other.add_run(reason_text)

    doc.add_paragraph()
    doc.add_paragraph(
        "4. В настоящее время: ☐ не работаю   ☐ работаю"
    )
    doc.add_paragraph(
        "На иждивении находятся _____ нетрудоспособных членов семьи (при отсутствии — «нет»)."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "5. Предупреждён(а) о необходимости извещать территориальный орган Фонда пенсионного "
        "и социального страхования РФ об обстоятельствах, влекущих изменение размера пенсии "
        "(ст. 26, 28 ФЗ № 400-ФЗ «О страховых пенсиях»)."
    )

    doc.add_paragraph()
    doc.add_paragraph("6. К заявлению прилагаю документы: _________________________________________")

    doc.add_paragraph()
    date_str = datetime.now().strftime("%d.%m.%Y")
    p_sign = doc.add_paragraph()
    p_sign.add_run(f"Дата заполнения заявления: {date_str}" + " " * 10 + "Подпись: _______________")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Документ подготовлен сервисом «Твой Голос» на основе официальной формы "
        "(Приложение N 2 к Административному регламенту, утв. Постановлением Правления "
        "ПФ РФ от 23.01.2019 N 16п, ред. от 23.09.2020). Персональные данные (ФИО, СНИЛС, "
        "адрес, вид пенсии) не заполняются автоматически — заполните их самостоятельно "
        "перед подачей. Сервис не является юридической фирмой и не несёт ответственности "
        "за содержание поданного документа."
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_edv_form(region_name: str, subcategory: str) -> bytes:
    """
    Генерирует бланк «Заявление о назначении ежемесячной денежной выплаты»
    (Приложение N 1, Постановление Правления ПФ РФ от 19.08.2019 N 414п).

    Из полной официальной формы намеренно НЕ включены: таблица членов семьи
    Героя СССР/РФ/кавалера ордена Славы (п.6 — узкий частный случай, не
    относящийся к нашей аудитории «инвалидность»/«ветеран боевых действий»),
    секретный вопрос/код для телефонной идентификации и способ уведомления
    (п.9-12 — опциональные разделы, заполняются по желанию самого заявителя,
    не связаны с сутью обращения). Все обязательные разделы (личные данные,
    представитель, категория и правовое основание, адресация в терр. орган,
    предупреждение об обязанностях, подпись) — сохранены.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"В Отделение Фонда пенсионного и социального страхования\nРоссийской Федерации по {to_dative(region_name)}")
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ЗАЯВЛЕНИЕ\nО НАЗНАЧЕНИИ ЕЖЕМЕСЯЧНОЙ ДЕНЕЖНОЙ ВЫПЛАТЫ")
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph(
        "1. Я, _________________________________________________________________________,"
    )
    doc.add_paragraph("(фамилия, имя, отчество (при наличии))").italic = True
    doc.add_paragraph(
        "Фамилия, которая была при рождении: _______________________________________"
    )
    _add_blank_line(doc, "СНИЛС", 30)
    _add_blank_line(doc, "Гражданство", 40)
    _add_blank_line(doc, "Адрес места жительства", 55)
    _add_blank_line(doc, "Документ, удостоверяющий личность (серия, номер, дата выдачи, кем выдан)", 45)
    _add_blank_line(doc, "Дата и место рождения", 45)
    doc.add_paragraph("Пол: ☐ муж.   ☐ жен.")

    doc.add_paragraph()
    doc.add_paragraph(
        "2. Представитель (если заявление подаётся представителем) — заполняется при "
        "необходимости, иначе оставить пустым: _______________________________________"
    )

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.add_run("4. Прошу назначить мне ежемесячную денежную выплату по категории ")
    p4.add_run("_________________________________ ").bold = False
    note = p4.add_run("(указать свою категорию самостоятельно)")
    note.italic = True
    note.font.size = Pt(9)

    p4b = doc.add_paragraph()
    legal_basis = EDV_LEGAL_BASIS.get(subcategory, "")
    if legal_basis:
        p4b.add_run("в соответствии с ")
        p4b.add_run(legal_basis).bold = True
        p4b.add_run(".")
    else:
        p4b.add_run("в соответствии с Федеральным законом ______________________________________ "
                     "(указать самостоятельно).")

    doc.add_paragraph()
    doc.add_paragraph(
        "5. Прошу направить выплатное дело получателя ежемесячной денежной выплаты в "
        "_______________________________________ (наименование территориального органа, "
        "если отличается от указанного в шапке заявления)."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "7. Документ, подтверждающий право на назначение ежемесячной денежной выплаты "
        "(наименование, серия, номер, дата выдачи, кем выдан): "
        "_______________________________________________________________________"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "8. Предупреждён(а) о необходимости безотлагательно извещать территориальный орган "
        "Фонда пенсионного и социального страхования РФ об обстоятельствах, влияющих на "
        "изменение размера ежемесячной денежной выплаты либо влекущих прекращение её выплаты. "
        "В случае получения излишних сумм социальных выплат в связи с несообщением о "
        "наступлении вышеуказанных обстоятельств обязуюсь возместить причинённый ущерб."
    )

    doc.add_paragraph()
    doc.add_paragraph("9. Контактный телефон: _______________________________")

    doc.add_paragraph()
    date_str = datetime.now().strftime("%d.%m.%Y")
    p_sign = doc.add_paragraph()
    p_sign.add_run(f"Дата заполнения заявления: {date_str}" + " " * 10 + "Подпись: _______________")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Документ подготовлен сервисом «Твой Голос» на основе официальной формы "
        "(Приложение N 1 к Административному регламенту, утв. Постановлением Правления "
        "ПФ РФ от 19.08.2019 N 414п, ред. от 23.09.2020). Из полной формы не включены "
        "опциональные разделы (способ уведомления, секретный код для телефонной "
        "идентификации, таблица членов семьи Героя СССР/РФ) — заполните их в оригинальном "
        "бланке при необходимости. Персональные данные не заполняются автоматически — "
        "заполните их самостоятельно перед подачей. Сервис не является юридической фирмой "
        "и не несёт ответственности за содержание поданного документа."
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_document(category: str, subcategory: str, region_name: str, reason_text: str) -> bytes:
    key = (category, subcategory)
    if key not in SUPPORTED_DOCUMENTS:
        raise ValueError(f"Document generation not supported for {key}")
    template_id = SUPPORTED_DOCUMENTS[key]
    if template_id == "pension_recalculation_form":
        return generate_pension_recalculation_form(region_name, reason_text)
    if template_id == "edv_form":
        return generate_edv_form(region_name, subcategory)
    raise ValueError(f"Unknown template_id {template_id}")
